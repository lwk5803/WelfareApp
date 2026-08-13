"""
document.py
------------
회원 등록 서류를 .docx(워드) 파일로 만듭니다. 기존에 복지관에서 종이로 받던
"이용신청서" + "개인정보 등 수집 및 이용 동의서"를 온라인 등록 절차에 맞게
정리한 것으로, A4 기준 2페이지(1페이지: 회원 등록 신청서, 2페이지: 개인정보
동의서)로 구성됩니다.
"""
import base64
import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# 기관명입니다. 다른 기관에서 이 앱을 쓰게 되면 여기 한 줄만 바꾸면 됩니다.
ORG_NAME = "구립응암노인복지관"

_HEADER_SHADE = "F2CFE3"  # 종이 서식의 분홍색 헤더 칸과 비슷한 색
_TITLE_COLOR = RGBColor(0xE0, 0x5A, 0x00)


def _shade_cell(cell, hex_color: str):
    """표 칸에 배경색을 칠합니다 (python-docx는 이 기능을 기본 API로 제공하지 않아
    문서 내부 XML을 직접 건드립니다)."""
    shd = cell._tc.get_or_add_tcPr().makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): hex_color,
    })
    cell._tc.get_or_add_tcPr().append(shd)


def _set_cell_text(cell, text: str, bold=False, size=10, shade: str | None = None, align_center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if shade:
        _shade_cell(cell, shade)


def _consent_mark(value: str) -> str:
    """저장된 동의 값("동의함"/"동의안함")을 체크박스 표시가 있는 문구로 바꿉니다."""
    agreed = value == "동의함"
    return f"{'☑' if agreed else '☐'} 동의함    {'☐' if agreed else '☑'} 동의하지 않음"


def _add_title(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = _TITLE_COLOR


def _add_org_footer(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(ORG_NAME)
    run.bold = True
    run.font.size = Pt(16)


def _build_page1(doc: Document, client: dict):
    _add_title(doc, "회원 등록 신청서")
    doc.add_paragraph()

    # --- 인적사항 + 사진 ---
    table = doc.add_table(rows=4, cols=5)
    table.style = "Table Grid"
    table.columns[0].width = Cm(3.2)
    photo_cell = table.cell(0, 0).merge(table.cell(3, 0))
    photo_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    photo_data = client.get("photo_data")
    if photo_data:
        try:
            img_stream = io.BytesIO(base64.b64decode(photo_data))
            photo_cell.paragraphs[0].add_run().add_picture(img_stream, width=Cm(2.8))
        except Exception:
            photo_cell.paragraphs[0].add_run("사진 오류")
    else:
        run = photo_cell.paragraphs[0].add_run("사진 없음")
        run.font.size = Pt(9)

    rows_info = [
        ("성명", client.get("name", ""), "회원번호", client.get("member_no", "") or "-"),
        ("생년월일", client.get("birth_date", ""), "성별", client.get("gender", "")),
        ("연락처", client.get("phone", ""), "상담자", client.get("counselor", "") or "-"),
        ("주소", client.get("address", ""), "", ""),
    ]
    for r, (label1, val1, label2, val2) in enumerate(rows_info):
        _set_cell_text(table.cell(r, 1), label1, bold=True, shade=_HEADER_SHADE)
        if label1 == "주소":
            # 주소는 길어서 오른쪽 칸까지 합쳐서 넓게 씁니다.
            wide = table.cell(r, 2).merge(table.cell(r, 3)).merge(table.cell(r, 4))
            _set_cell_text(wide, val1)
        else:
            _set_cell_text(table.cell(r, 2), val1)
            _set_cell_text(table.cell(r, 3), label2, bold=True, shade=_HEADER_SHADE)
            _set_cell_text(table.cell(r, 4), val2)

    # --- 비상연락망 ---
    doc.add_paragraph()
    emg_table = doc.add_table(rows=1, cols=4)
    emg_table.style = "Table Grid"
    _set_cell_text(emg_table.cell(0, 0), "비상연락망", bold=True, shade=_HEADER_SHADE)
    emg_text = (
        f"성명: {client.get('emergency_contact_name') or '-'}   "
        f"관계: {client.get('emergency_contact_relation') or '-'}   "
        f"연락처: {client.get('emergency_contact_phone') or '-'}"
    )
    wide = emg_table.cell(0, 1).merge(emg_table.cell(0, 2)).merge(emg_table.cell(0, 3))
    _set_cell_text(wide, emg_text)

    # --- 상담/생활 정보 ---
    doc.add_paragraph()
    illness = client.get("has_illness", "없다")
    illness_line = illness if illness != "있다" else f"있다 ({client.get('illness_type') or '미기재'})"
    career = client.get("has_career", "없다")
    career_line = career if career != "있다" else f"있다 ({client.get('career_type') or '미기재'})"
    disability = client.get("has_disability", "아니오")
    disability_line = disability if disability != "예" else f"예 ({client.get('disability_type') or '미기재'})"

    detail_rows = [
        ("상담일자", (client.get("created_at") or "")[:10] or "-", "생활구분", client.get("welfare_type", "")),
        ("가입경로", client.get("join_route") or "-", "동거형태", client.get("household_types") or "-"),
        ("질병유무", illness_line, "장애유무", disability_line),
        ("경력(전 직업)", career_line, "", ""),
        ("관심있는 서비스", "", "", ""),
        ("관심있는 것은?", "", "", ""),
        ("상담결과", "", "", ""),
    ]
    dt = doc.add_table(rows=len(detail_rows), cols=4)
    dt.style = "Table Grid"
    for r, (label1, val1, label2, val2) in enumerate(detail_rows):
        _set_cell_text(dt.cell(r, 0), label1, bold=True, shade=_HEADER_SHADE)
        if label2:
            _set_cell_text(dt.cell(r, 1), val1)
            _set_cell_text(dt.cell(r, 2), label2, bold=True, shade=_HEADER_SHADE)
            _set_cell_text(dt.cell(r, 3), val2)
        else:
            wide = dt.cell(r, 1).merge(dt.cell(r, 2)).merge(dt.cell(r, 3))
            _set_cell_text(wide, val1)

    doc.add_paragraph()
    _add_org_footer(doc)


def _add_consent_table(doc: Document, rows: list[tuple[str, str]], consent_row_label: str, consent_value: str):
    """개인정보 동의서의 표 하나(항목/설명 표 + 마지막 동의여부 행)를 만듭니다."""
    table = doc.add_table(rows=len(rows) + 1, cols=2)
    table.style = "Table Grid"
    table.columns[0].width = Cm(4.2)
    for r, (label, value) in enumerate(rows):
        _set_cell_text(table.cell(r, 0), label, bold=True, shade=_HEADER_SHADE, size=9)
        _set_cell_text(table.cell(r, 1), value, size=9)
    _set_cell_text(table.cell(len(rows), 0), consent_row_label, bold=True, shade=_HEADER_SHADE, size=9)
    _set_cell_text(table.cell(len(rows), 1), consent_value, bold=True, size=9)


def _build_page2(doc: Document, client: dict):
    _add_title(doc, "개인정보 등 수집 및 이용 동의서")
    p = doc.add_paragraph(
        f"{ORG_NAME} 이용 서비스 신청과 관련하여 개인정보보호법 제15조 및 제22조에 따라 귀하의 "
        "개인정보 및 민감정보를 수집·이용하고자 하오니 아래 사항을 숙지하신 후 동의여부를 확인하여 주시기 바랍니다."
    )
    for run in p.runs:
        run.font.size = Pt(9)
    doc.add_paragraph()

    h1 = doc.add_paragraph()
    h1.add_run("□ 개인정보 및 민감정보 수집·이용 동의").bold = True
    _add_consent_table(
        doc,
        [
            ("수집·이용 목적", f"{ORG_NAME} 회원가입(사회복지시설정보시스템)과 프로그램 참여 및 서비스 제공·관리"),
            ("수집하는 기본 개인정보 항목", "성명, 성별, 생년월일, 주소, 연락처 등"),
            ("수집하는 민감정보 항목", "건강정보(질병·장애 유무), 소득 및 재산정보, 보호구분 등"),
            ("보유·이용 기간", "이용신청서 제출 시부터 회원 자격 상실 시까지"),
            ("동의거부 권리 및 제한사항", "귀하는 동의를 거부할 권리가 있으며, 거부 시 회원가입과 서비스 이용이 제한될 수 있습니다."),
        ],
        "개인정보/민감정보 수집·이용 동의 여부",
        _consent_mark(client.get("consent_personal", "")) + " / " + _consent_mark(client.get("consent_sensitive", "")),
    )
    doc.add_paragraph()

    h2 = doc.add_paragraph()
    h2.add_run("□ 개인정보 제3자 제공 동의").bold = True
    _add_consent_table(
        doc,
        [
            ("제공받는 자", "해당 지자체 행정기관(시청, 구청, 주민센터 등) 및 사회복지서비스 제공기관"),
            ("이용 목적", "관계법령에 따른 제출·요구, 서비스 의뢰, 사업소개·홍보 등 사회복지서비스 이용에 필요한 경우"),
            ("제공되는 항목", "성명, 성별, 생년월일, 주소, 연락처, 건강정보, 소득 및 재산정보, 보호구분 등"),
            ("이용·보유 기간", "해당 서비스 지원과 행정사무가 제공되는 기간"),
            ("동의거부 권리 및 제한사항", "귀하는 동의를 거부할 권리가 있으며, 거부 시 회원가입과 서비스 이용이 제한될 수 있습니다."),
        ],
        "개인정보 제3자 제공 동의 여부",
        _consent_mark(client.get("consent_third_party", "")),
    )
    doc.add_paragraph()

    h3 = doc.add_paragraph()
    h3.add_run("□ 초상권 사용 및 이용 동의").bold = True
    _add_consent_table(
        doc,
        [
            ("사용·이용 목적", f"{ORG_NAME} 온·오프라인 사업소개 및 홍보, 기록자료 활용"),
            ("사용·이용 항목", "서비스 이용사진 및 운영사진 등 초상사진저작물"),
            ("동의거부 권리 및 제한사항", "귀하는 동의를 거부할 권리가 있으며, 거부하셔도 서비스 이용에는 제한이 없습니다."),
        ],
        "초상권 사용·이용 동의 여부",
        _consent_mark(client.get("consent_portrait", "")),
    )

    doc.add_paragraph()
    signed_at = (client.get("consent_signed_at") or "")[:10]
    p = doc.add_paragraph(f"동의 확인일: {signed_at or '-'}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2 = doc.add_paragraph(f"신청인: {client.get('name', '')}  (서명 또는 인)")
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph()
    _add_org_footer(doc)


def build_registration_document(client: dict) -> bytes:
    """회원 한 명의 정보(dict, get_client() 결과)로 .docx 문서(bytes)를 만듭니다."""
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    for margin in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, margin, Cm(1.8))

    _build_page1(doc, client)
    doc.add_page_break()
    _build_page2(doc, client)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
